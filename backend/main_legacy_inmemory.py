"""
Production Backend - AI Document Classification
SQLite Database + Text Extraction + AI Features
"""
import os
import re
import json
import uuid
import hashlib
from datetime import datetime, timedelta
from typing import List, Optional
from pathlib import Path

from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, Form, Query, status
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from sqlmodel import SQLModel, Session, create_engine, select, Field
from pydantic import BaseModel

import PyPDF2
from docx import Document as DocxDocument
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer

# ============================================================================
# DATABASE CONFIGURATION
# ============================================================================

DATABASE_URL = "sqlite:///documents.db"
engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})

def create_db_and_tables():
    SQLModel.metadata.create_all(engine)

def get_session():
    with Session(engine) as session:
        yield session

# ============================================================================
# DATABASE MODELS
# ============================================================================

class User(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    username: str = Field(unique=True, index=True)
    password: str
    full_name: str
    email: str
    role: str = "user"  # admin, hr, finance, user
    created_at: datetime = Field(default_factory=datetime.utcnow)

class Document(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    filename: str = Field(index=True)
    original_filename: str
    category: str = "General"
    size: int
    path: str
    author: Optional[str] = None
    title: Optional[str] = None
    summary: Optional[str] = None
    text_content: Optional[str] = None
    entities: Optional[str] = None  # JSON string
    mime_type: str
    uploader_id: int = Field(foreign_key="user.id")
    upload_date: datetime = Field(default_factory=datetime.utcnow)
    last_accessed: Optional[datetime] = None

class AccessLog(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    user_id: int = Field(foreign_key="user.id")
    document_id: Optional[int] = Field(foreign_key="document.id", default=None)
    action: str  # upload, view, download, delete
    timestamp: datetime = Field(default_factory=datetime.utcnow)

# ============================================================================
# PYDANTIC MODELS (FOR API)
# ============================================================================

class UserLogin(BaseModel):
    username: str
    password: str

class UserResponse(BaseModel):
    id: int
    username: str
    full_name: str
    email: str
    role: str

class DocumentResponse(BaseModel):
    id: int
    name: str
    original_filename: str
    category: str
    size: int
    author: Optional[str]
    mime_type: str
    upload_date: datetime
    last_accessed: Optional[datetime]

class DocumentDetail(DocumentResponse):
    title: Optional[str]
    summary: Optional[str]
    entities: Optional[dict]
    text_content: Optional[str]

class TokenResponse(BaseModel):
    access_token: str
    token_type: str
    user: UserResponse

# ============================================================================
# FASTAPI APP
# ============================================================================

app = FastAPI(title="AI Document Classification API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ============================================================================
# AUTHENTICATION
# ============================================================================

# In-memory token store (in production, use JWT)
tokens_store = {}

def get_password_hash(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return get_password_hash(plain_password) == hashed_password

def get_current_user(token: str = None, session: Session = Depends(get_session)) -> User:
    if not token:
        raise HTTPException(status_code=401, detail="Missing token")
    
    user_id = tokens_store.get(token)
    if not user_id:
        raise HTTPException(status_code=401, detail="Invalid token")
    
    user = session.get(User, user_id)
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    
    return user

# ============================================================================
# TEXT EXTRACTION
# ============================================================================

def extract_text_from_pdf(file_path: str, max_pages: int = 10) -> str:
    """Extract text from PDF (first 10 pages)"""
    try:
        text = []
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for i, page in enumerate(reader.pages[:max_pages]):
                text.append(page.extract_text())
        return "\n".join(text)
    except Exception as e:
        return ""

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX"""
    try:
        doc = DocxDocument(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        return ""

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        return ""

def extract_text(file_path: str) -> str:
    """Route to correct extraction function"""
    ext = Path(file_path).suffix.lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.txt':
        return extract_text_from_txt(file_path)
    
    return ""

# ============================================================================
# METADATA EXTRACTION
# ============================================================================

def extract_metadata(text: str, filename: str) -> dict:
    """Extract title, author, date from text"""
    
    # Title: first non-empty line
    title = None
    for line in text.split('\n')[:10]:
        if line.strip() and len(line.strip()) < 150:
            title = line.strip()[:100]
            break
    if not title:
        title = filename
    
    # Author: look for common patterns
    author = None
    patterns = [
        r'(?:Author|By|From):\s*([^\n]+)',
        r'(?:Author|By):\s*([^\n]+)',
    ]
    for pattern in patterns:
        match = re.search(pattern, text[:1000], re.IGNORECASE)
        if match:
            author = match.group(1).strip()[:100]
            break
    
    # Date: look for date patterns
    date_pattern = r'\b(?:20\d{2}|19\d{2})[-/](?:0?[1-9]|1[0-2])[-/](?:0?[1-9]|[12]\d|3[01])\b'
    date_match = re.search(date_pattern, text)
    date = date_match.group(0) if date_match else None
    
    return {
        "title": title,
        "author": author,
        "date": date,
    }

# ============================================================================
# ENTITY EXTRACTION
# ============================================================================

def extract_entities(text: str) -> dict:
    """Extract emails, amounts, phone numbers"""
    
    emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    amounts = re.findall(r'\$[\d,]+(?:\.\d{2})?', text)
    phone_numbers = re.findall(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
    
    return {
        "emails": list(set(emails)),
        "amounts": list(set(amounts)),
        "phone_numbers": list(set(phone_numbers)),
    }

# ============================================================================
# SUMMARIZATION (TF-IDF EXTRACTIVE)
# ============================================================================

def summarize_text(text: str, num_sentences: int = 3) -> str:
    """Extract top N sentences by TF-IDF score"""
    try:
        # Split into sentences
        sentences = re.split(r'(?<=[.!?])\s+', text.strip())[:20]
        
        if len(sentences) <= num_sentences:
            return " ".join(sentences)
        
        # TF-IDF vectorization
        vectorizer = TfidfVectorizer(max_features=100, stop_words='english')
        tfidf = vectorizer.fit_transform(sentences)
        
        # Sum TF-IDF scores for each sentence
        scores = np.asarray(tfidf.sum(axis=1)).flatten()
        
        # Get top N sentence indices
        top_indices = np.argsort(scores)[-num_sentences:]
        top_indices = sorted(top_indices)
        
        # Return summary in original order
        return " ".join([sentences[i] for i in top_indices])
    except Exception as e:
        return text[:500]

# ============================================================================
# DOCUMENT CLASSIFICATION
# ============================================================================

def classify_category(text: str, filename: str) -> str:
    """Classify document into category"""
    
    combined_text = (text + " " + filename).lower()[:2000]
    
    categories = {
        "Finance": ["invoice", "receipt", "payment", "financial", "budget", "balance", "expense", "revenue"],
        "HR": ["employee", "recruitment", "salary", "benefits", "hr", "human resources", "hiring", "performance"],
        "Legal": ["contract", "agreement", "legal", "terms", "conditions", "law", "attorney", "court"],
        "Contracts": ["contract", "agreement", "terms", "conditions", "signature", "parties"],
        "Technical Reports": ["technical", "report", "system", "software", "development", "architecture"],
    }
    
    for category, keywords in categories.items():
        for keyword in keywords:
            if keyword in combined_text:
                return category
    
    return "General"

# ============================================================================
# SEMANTIC SEARCH (FALLBACK TO TEXT SEARCH)
# ============================================================================

def text_search(query: str, documents: List[Document]) -> List[tuple]:
    """Simple text-based search"""
    results = []
    query_lower = query.lower()
    
    for doc in documents:
        if not doc.text_content:
            continue
        
        content_lower = doc.text_content.lower()
        score = 0
        
        # Score based on term frequency
        for term in query_lower.split():
            score += content_lower.count(term)
        
        if score > 0:
            results.append((doc, score / 100.0))  # Normalize
    
    return sorted(results, key=lambda x: x[1], reverse=True)

# ============================================================================
# STARTUP
# ============================================================================

@app.on_event("startup")
def on_startup():
    create_db_and_tables()
    
    # Create default users if they don't exist
    with Session(engine) as session:
        for username, (password, full_name, email, role) in [
            ("admin", ("admin123", "Administrator", "admin@example.com", "admin")),
            ("hr", ("hr123", "HR Manager", "hr@example.com", "hr")),
            ("finance", ("finance123", "Finance Manager", "finance@example.com", "finance")),
        ].items():
            existing = session.exec(select(User).where(User.username == username)).first()
            if not existing:
                user = User(
                    username=username,
                    password=get_password_hash(password),
                    full_name=full_name,
                    email=email,
                    role=role,
                )
                session.add(user)
        session.commit()

# ============================================================================
# AUTH ENDPOINTS
# ============================================================================

@app.post("/auth/login", response_model=TokenResponse)
def login(credentials: UserLogin, session: Session = Depends(get_session)):
    """Login with username/password"""
    user = session.exec(select(User).where(User.username == credentials.username)).first()
    
    if not user or not verify_password(credentials.password, user.password):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    token = str(uuid.uuid4())
    tokens_store[token] = user.id
    
    return TokenResponse(
        access_token=token,
        token_type="bearer",
        user=UserResponse(
            id=user.id,
            username=user.username,
            full_name=user.full_name,
            email=user.email,
            role=user.role,
        ),
    )

@app.post("/auth/logout")
def logout(token: str = None):
    """Logout (invalidate token)"""
    if token and token in tokens_store:
        del tokens_store[token]
    return {"message": "Logged out"}

@app.get("/auth/me", response_model=UserResponse)
def get_me(current_user: User = Depends(get_current_user)):
    """Get current user info"""
    return UserResponse(
        id=current_user.id,
        username=current_user.username,
        full_name=current_user.full_name,
        email=current_user.email,
        role=current_user.role,
    )

# ============================================================================
# DOCUMENT ENDPOINTS
# ============================================================================

@app.post("/documents/upload")
def upload_document(
    file: UploadFile = File(...),
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    """Upload and process document"""
    
    # Create uploads directory if needed
    os.makedirs("uploads", exist_ok=True)
    
    # Save file
    file_id = str(uuid.uuid4())
    ext = Path(file.filename).suffix
    saved_filename = f"{file_id}{ext}"
    file_path = os.path.join("uploads", saved_filename)
    
    with open(file_path, "wb") as f:
        f.write(file.file.read())
    
    file_size = os.path.getsize(file_path)
    
    # Extract text
    text_content = extract_text(file_path)
    if not text_content:
        text_content = "[Unable to extract text]"
    
    # Extract metadata
    metadata = extract_metadata(text_content, file.filename)
    
    # Extract entities
    entities = extract_entities(text_content)
    
    # Classify category
    category = classify_category(text_content, file.filename)
    
    # Summarize
    summary = summarize_text(text_content, num_sentences=3)
    
    # Create document record
    doc = Document(
        filename=saved_filename,
        original_filename=file.filename,
        category=category,
        size=file_size,
        path=file_path,
        author=metadata.get("author") or current_user.full_name,
        title=metadata.get("title"),
        summary=summary,
        text_content=text_content[:5000],  # Store first 5000 chars
        entities=json.dumps(entities),
        mime_type=file.content_type or "application/octet-stream",
        uploader_id=current_user.id,
    )
    
    session.add(doc)
    session.commit()
    session.refresh(doc)
    
    # Log access
    log = AccessLog(user_id=current_user.id, document_id=doc.id, action="upload")
    session.add(log)
    session.commit()
    
    return DocumentResponse(
        id=doc.id,
        name=doc.filename,
        original_filename=doc.original_filename,
        category=doc.category,
        size=doc.size,
        author=doc.author,
        mime_type=doc.mime_type,
        upload_date=doc.upload_date,
        last_accessed=doc.last_accessed,
    )

@app.get("/documents/", response_model=List[DocumentResponse])
def list_documents(
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    """List documents (filtered by role)"""
    
    query = select(Document)
    
    # Role-based filtering
    if current_user.role == "hr":
        query = query.where(Document.category == "HR")
    elif current_user.role == "finance":
        query = query.where(Document.category == "Finance")
    # admin sees all
    
    documents = session.exec(query).all()
    
    return [
        DocumentResponse(
            id=doc.id,
            name=doc.filename,
            original_filename=doc.original_filename,
            category=doc.category,
            size=doc.size,
            author=doc.author,
            mime_type=doc.mime_type,
            upload_date=doc.upload_date,
            last_accessed=doc.last_accessed,
        )
        for doc in documents
    ]

@app.get("/documents/{doc_id}", response_model=DocumentDetail)
def get_document(
    doc_id: int,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    """Get document details"""
    
    doc = session.get(Document, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Role-based access check
    if current_user.role == "hr" and doc.category != "HR":
        raise HTTPException(status_code=403, detail="Access denied")
    elif current_user.role == "finance" and doc.category != "Finance":
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Update last accessed
    doc.last_accessed = datetime.utcnow()
    session.add(doc)
    session.commit()
    
    # Log access
    log = AccessLog(user_id=current_user.id, document_id=doc.id, action="view")
    session.add(log)
    session.commit()
    
    entities = json.loads(doc.entities) if doc.entities else {}
    
    return DocumentDetail(
        id=doc.id,
        name=doc.filename,
        original_filename=doc.original_filename,
        category=doc.category,
        size=doc.size,
        author=doc.author,
        mime_type=doc.mime_type,
        upload_date=doc.upload_date,
        last_accessed=doc.last_accessed,
        title=doc.title,
        summary=doc.summary,
        entities=entities,
        text_content=doc.text_content,
    )

@app.get("/documents/{doc_id}/download")
def download_document(
    doc_id: int,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    """Download document file"""
    
    doc = session.get(Document, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Role-based access check
    if current_user.role == "hr" and doc.category != "HR":
        raise HTTPException(status_code=403, detail="Access denied")
    elif current_user.role == "finance" and doc.category != "Finance":
        raise HTTPException(status_code=403, detail="Access denied")
    
    if not os.path.exists(doc.path):
        raise HTTPException(status_code=404, detail="File not found")
    
    # Log access
    log = AccessLog(user_id=current_user.id, document_id=doc.id, action="download")
    session.add(log)
    session.commit()
    
    return FileResponse(doc.path, filename=doc.original_filename)

@app.delete("/documents/{doc_id}")
def delete_document(
    doc_id: int,
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    """Delete document"""
    
    doc = session.get(Document, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Only uploader or admin can delete
    if doc.uploader_id != current_user.id and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Delete file
    if os.path.exists(doc.path):
        os.remove(doc.path)
    
    # Log access
    log = AccessLog(user_id=current_user.id, document_id=doc.id, action="delete")
    session.add(log)
    session.commit()
    
    session.delete(doc)
    session.commit()
    
    return {"message": "Document deleted"}

# ============================================================================
# SEARCH ENDPOINTS
# ============================================================================

@app.get("/search/")
def search(
    q: str = Query(..., min_length=1),
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    """Text search documents"""
    
    # Get all documents (will filter by role)
    query = select(Document)
    if current_user.role == "hr":
        query = query.where(Document.category == "HR")
    elif current_user.role == "finance":
        query = query.where(Document.category == "Finance")
    
    documents = session.exec(query).all()
    
    # Text search
    results = text_search(q, documents)
    
    return [
        DocumentResponse(
            id=doc.id,
            name=doc.filename,
            original_filename=doc.original_filename,
            category=doc.category,
            size=doc.size,
            author=doc.author,
            mime_type=doc.mime_type,
            upload_date=doc.upload_date,
            last_accessed=doc.last_accessed,
        )
        for doc, score in results[:10]
    ]

# ============================================================================
# ACCESS LOG ENDPOINTS (ADMIN ONLY)
# ============================================================================

@app.get("/logs/")
def get_logs(
    session: Session = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    """Get access logs (admin only)"""
    
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin only")
    
    logs = session.exec(select(AccessLog).order_by(AccessLog.timestamp.desc())).all()
    
    return [
        {
            "id": log.id,
            "user_id": log.user_id,
            "document_id": log.document_id,
            "action": log.action,
            "timestamp": log.timestamp.isoformat(),
        }
        for log in logs
    ]

# ============================================================================
# HEALTH CHECK
# ============================================================================

@app.get("/health")
def health():
    """Health check endpoint"""
    return {"status": "ok"}
